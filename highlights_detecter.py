import docx
import json

class entity:
    count = 0
    def __init__(self, text, classification):
        self.id = entity.count
        entity.count += 1
        self.text = text
        self.classification = classification
        self.EDSS = 0.0 

    def to_dict(self):
        return {
            "id": self.id,
            "text": self.text,
            "classification": self.classification,
            "EDSS": self.EDSS
        }

def file_open(filename):
    try:
        document = docx.Document(filename + ".docx")
        return document
    except:
        NameError = "File not found"
        print(NameError)
        return None

def get_edss(id, text):
    return 0.0

def get_classification(shading_val):
    match shading_val:
        case None:
            return None
        case "FCFCFC":
            return "Высшие корковые"
        case "BDD6EE":
            return "Зрительная"
        case "C5E0B3":
            return "Стволовая"
        case "B4C6E7":
            return "Пирамидная"
        case "F7CAAC":
            return "Мозжечковая"
        case "FFF2CC":
            return "Чувствительная"
        case "DBDBDB":
            return "Тазовые"
        case "F2F2F2":
            return "Амб. индекс."

document = docx.Document("file_ref.docx")
stored_lib = []
for paragraph in document.paragraphs[1:]: 
    existing_entry = None
    for run in paragraph.runs:
        r_element = run._element
        shading = r_element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        shading_val = shading.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill") if shading is not None else None
        if shading_val is not None:
            existing_entry = next((entry for entry in stored_lib if entry.classification == get_classification(shading_val)), None)
            if existing_entry:
                existing_entry.text += " " + run.text
            else:
                entry = entity(run.text, get_classification(shading_val))
                if entry.classification == "null":
                    pass
                else:
                    stored_lib.append(entry)

print(stored_lib)
with open("marked.json", "w") as json_file:
    json.dump([entry.to_dict() for entry in stored_lib], json_file, ensure_ascii=False, indent=4)
